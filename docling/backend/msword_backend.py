import logging
import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Final, Optional, Union

from docling_core.types.doc import (
    ContentLayer,
    DocItemLabel,
    DoclingDocument,
    DocumentOrigin,
    GroupLabel,
    ImageRef,
    ListGroup,
    NodeItem,
    RefItem,
    RichTableCell,
    TableCell,
    TableData,
    TableItem,
)
from docling_core.types.doc.document import Formatting, Script
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tc
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.table import Table, _Cell
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree
from PIL import Image, UnidentifiedImageError
from pydantic import AnyUrl
from typing_extensions import override

from docling.backend.abstract_backend import DeclarativeDocumentBackend
from docling.backend.docx.drawingml.utils import (
    get_docx_to_pdf_converter,
    get_pil_from_dml_docx,
)
from docling.backend.docx.latex.omml import oMath2Latex
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import InputDocument

_log = logging.getLogger(__name__)


class MsWordDocumentBackend(DeclarativeDocumentBackend):
    """
    Microsoft Word 文档（DOCX 格式）解析后端。
    
    该类负责将 DOCX 文件解析为 DoclingDocument 结构，处理文档中的各种元素，
    包括文本、表格、图片、公式、列表、文本框、页眉页脚等。
    """
    _BLIP_NAMESPACES: Final = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "v": "urn:schemas-microsoft-com:vml",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        "w10": "urn:schemas-microsoft-com:office:word",
        "a14": "http://schemas.microsoft.com/office/drawing/2010/main",
    }
    """
    Word 文档中使用的 XML 命名空间映射。
    
    这些命名空间用于解析 DOCX 文件中的各种元素，包括：
    - a: DrawingML 主命名空间
    - r: Office 文档关系命名空间
    - w: WordprocessingML 主命名空间
    - wp: Wordprocessing Drawing 命名空间
    - mc: 标记兼容性命名空间
    - v: VML (Vector Markup Language) 命名空间
    - wps: Wordprocessing Shape 命名空间
    - w10: Office Word 命名空间
    - a14: Office 2010 Drawing 命名空间
    """

    @override
    def __init__(
        self, in_doc: "InputDocument", path_or_stream: Union[BytesIO, Path]
    ) -> None:
        """
        初始化 MsWordDocumentBackend 实例。
        
        Args:
            in_doc: 输入文档对象
            path_or_stream: DOCX 文件路径或字节流
        """
        super().__init__(in_doc, path_or_stream)
        self.XML_KEY = (
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
        )
        self.xml_namespaces = {
            "w": "http://schemas.microsoft.com/office/word/2003/wordml"
        }
        self.blip_xpath_expr = etree.XPath(
            ".//a:blip", namespaces=MsWordDocumentBackend._BLIP_NAMESPACES
        )
        # self.initialise(path_or_stream)
        # Word file:
        self.path_or_stream: Union[BytesIO, Path] = path_or_stream
        self.valid: bool = False
        # Initialise the parents for the hierarchy
        self.max_levels: int = 10  # 文档层次结构的最大层级数
        self.level_at_new_list: Optional[int] = None  # 新列表开始时的层级
        self.parents: dict[int, Optional[NodeItem]] = {}  # 父级节点映射
        self.numbered_headers: dict[int, int] = {}  # 编号标题计数器
        self.equation_bookends: str = "<eq>{EQ}</eq>"  # 公式标记格式
        # Track processed textbox elements to avoid duplication
        self.processed_textbox_elements: list[int] = []  # 已处理的文本框元素ID列表
        self.docx_to_pdf_converter: Optional[Callable] = None  # DOCX到PDF转换器
        self.docx_to_pdf_converter_init = False  # 转换器初始化标志
        self.display_drawingml_warning = True  # 是否显示DrawingML警告

        # 初始化父级节点字典
        for i in range(-1, self.max_levels):
            self.parents[i] = None

        self.level = 0
        self.listIter = 0
        # Track list counters per numId and ilvl
        self.list_counters: dict[tuple[int, int], int] = {}  # 列表计数器 (numId, ilvl) -> count
        # Set starting content layer
        self.content_layer = ContentLayer.BODY  # 默认内容层为正文

        # 初始化处理历史记录
        self.history: dict[str, Any] = {
            "names": [None],     # 样式名称历史
            "levels": [None],    # 层级历史
            "numids": [None],    # 列表编号ID历史
            "indents": [None],   # 缩进级别历史
        }

        # 加载 DOCX 文件
        self.docx_obj = self.load_msword_file(
            path_or_stream=self.path_or_stream, document_hash=self.document_hash
        )
        if self.docx_obj:
            self.valid = True

    @override
    def is_valid(self) -> bool:
        """
        检查文档是否有效加载。
        
        Returns:
            bool: 如果文档有效则返回 True，否则返回 False
        """
        return self.valid

    @classmethod
    @override
    def supports_pagination(cls) -> bool:
        """
        检查是否支持分页。
        
        Returns:
            bool: DOCX 后端不支持分页，始终返回 False
        """
        return False

    @override
    def unload(self):
        """
        卸载文档资源，释放内存。
        """
        if isinstance(self.path_or_stream, BytesIO):
            self.path_or_stream.close()

        self.path_or_stream = None

    @classmethod
    @override
    def supported_formats(cls) -> set[InputFormat]:
        """
        返回支持的输入格式。
        
        Returns:
            set[InputFormat]: 支持 DOCX 格式
        """
        return {InputFormat.DOCX}

    @override
    def convert(self) -> DoclingDocument:
        """
        将 DOCX 文件解析为结构化文档模型。
        
        该方法是文档解析的核心，它遍历文档的所有元素，
        包括文本、表格、图片、公式等，并构建 DoclingDocument 对象。

        Returns:
            DoclingDocument: 解析后的文档对象
            
        Raises:
            RuntimeError: 当后端初始化失败时抛出异常
        """

        origin = DocumentOrigin(
            filename=self.file.name or "file",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            binary_hash=self.document_hash,
        )

        doc = DoclingDocument(name=self.file.stem or "file", origin=origin)
        if self.is_valid():
            assert self.docx_obj is not None
            # 遍历文档主体元素
            doc, _ = self._walk_linear(self.docx_obj.element.body, doc)
            # 添加页眉页脚
            self._add_header_footer(self.docx_obj, doc)

            return doc
        else:
            raise RuntimeError(
                f"Cannot convert doc with {self.document_hash} because the backend failed to init."
            )

    @staticmethod
    def load_msword_file(
        path_or_stream: Union[BytesIO, Path], document_hash: str
    ) -> DocxDocument:
        """
        加载 Microsoft Word 文档。
        
        Args:
            path_or_stream: DOCX 文件路径或字节流
            document_hash: 文档哈希值，用于错误信息
            
        Returns:
            DocxDocument: 加载的文档对象
            
        Raises:
            RuntimeError: 当文档加载失败时抛出异常
        """
        try:
            if isinstance(path_or_stream, BytesIO):
                return Document(path_or_stream)
            elif isinstance(path_or_stream, Path):
                return Document(str(path_or_stream))
            else:
                return None
        except Exception as e:
            raise RuntimeError(
                f"MsWordDocumentBackend could not load document with hash {document_hash}"
            ) from e

    def _update_history(
        self,
        name: str,
        level: Optional[int],
        numid: Optional[int],
        ilevel: Optional[int],
    ):
        """
        更新处理历史记录。
        
        Args:
            name: 样式名称
            level: 层级
            numid: 列表编号ID
            ilevel: 缩进级别
        """
        self.history["names"].append(name)
        self.history["levels"].append(level)

        self.history["numids"].append(numid)
        self.history["indents"].append(ilevel)

    def _prev_name(self) -> Optional[str]:
        """
        获取上一个处理元素的样式名称。
        
        Returns:
            Optional[str]: 上一个样式名称
        """
        return self.history["names"][-1]

    def _prev_level(self) -> Optional[int]:
        """
        获取上一个处理元素的层级。
        
        Returns:
            Optional[int]: 上一个层级
        """
        return self.history["levels"][-1]

    def _prev_numid(self) -> Optional[int]:
        """
        获取上一个处理元素的列表编号ID。
        
        Returns:
            Optional[int]: 上一个列表编号ID
        """
        return self.history["numids"][-1]

    def _prev_indent(self) -> Optional[int]:
        """
        获取上一个处理元素的缩进级别。
        
        Returns:
            Optional[int]: 上一个缩进级别
        """
        return self.history["indents"][-1]

    def _get_level(self) -> int:
        """
        获取当前可用的层级索引。
        
        返回父级节点字典中第一个值为 None 的键，表示当前可用的层级。
        
        Returns:
            int: 当前可用的层级索引
        """
        for k, v in self.parents.items():
            if k >= 0 and v is None:
                return k
        return 0

    def _walk_linear(
        self,
        body: BaseOxmlElement,
        doc: DoclingDocument,
        # parent:
    ) -> tuple[DoclingDocument, list[RefItem]]:
        """
        线性遍历文档元素。
        
        该方法是文档解析的核心，负责遍历文档中的所有元素，
        并根据元素类型调用相应的处理函数。
        
        Args:
            body: 文档主体元素
            doc: DoclingDocument 对象
            
        Returns:
            tuple[DoclingDocument, list[RefItem]]: 更新后的文档和添加的元素引用列表
        """
        # 存储已处理的元素引用
        added_elements = []
        
        # 遍历文档主体中的每一个元素
        for element in body:
            # 获取元素的标签名（去除命名空间前缀）
            tag_name = etree.QName(element).localname
            
            # 检查是否存在内联图像（blip元素）
            drawing_blip = self.blip_xpath_expr(element)
            
            # 查找所有绘图元素（用于处理DrawingML）
            drawingml_els = element.findall(
                ".//w:drawing", namespaces=MsWordDocumentBackend._BLIP_NAMESPACES
            )

            # 检查文本框内容 - 支持多种文本框格式
            # 仅当元素未被处理过时才进行处理，防止重复处理
            element_id = id(element)
            if element_id not in self.processed_textbox_elements:
                # 现代Word文本框的XPath表达式
                txbx_xpath = etree.XPath(
                    ".//w:txbxContent|.//v:textbox//w:p",
                    namespaces=MsWordDocumentBackend._BLIP_NAMESPACES,
                )
                textbox_elements = txbx_xpath(element)

                # 如果未找到现代文本框且当前元素是绘图或图片元素，
                # 则检查替代/旧版文本框格式
                if not textbox_elements and tag_name in ["drawing", "pict"]:
                    # DrawingML和VML格式中文本框的额外检查
                    alt_txbx_xpath = etree.XPath(
                        ".//wps:txbx//w:p|.//w10:wrap//w:p|.//a:p//a:t",
                        namespaces=MsWordDocumentBackend._BLIP_NAMESPACES,
                    )
                    textbox_elements = alt_txbx_xpath(element)

                    # 检查不在标准文本框中的形状文本
                    if not textbox_elements:
                        shape_text_xpath = etree.XPath(
                            ".//a:bodyPr/ancestor::*//a:t|.//a:txBody//a:t",
                            namespaces=MsWordDocumentBackend._BLIP_NAMESPACES,
                        )
                        shape_text_elements = shape_text_xpath(element)
                        if shape_text_elements:
                            # 从形状文本创建自定义文本元素
                            text_content = " ".join(
                                [t.text for t in shape_text_elements if t.text]
                            )
                            # 如果存在文本内容，则创建对应的文档元素
                            if text_content.strip():
                                _log.debug(f"Found shape text: {text_content[:50]}...")
                                # 创建段落式元素以便使用标准处理器处理
                                level = self._get_level()
                                shape_group = doc.add_group(
                                    label=GroupLabel.SECTION,
                                    parent=self.parents[level - 1],
                                    name="shape-text",
                                    content_layer=self.content_layer,
                                )
                                added_elements.append(shape_group.get_ref())
                                doc.add_text(
                                    label=DocItemLabel.TEXT,
                                    parent=shape_group,
                                    text=text_content,
                                    content_layer=self.content_layer,
                                )

                # 如果找到了文本框元素，则处理这些内容
                if textbox_elements:
                    # 标记父元素已被处理
                    self.processed_textbox_elements.append(element_id)
                    # 同时标记所有找到的文本框元素已被处理
                    for tb_element in textbox_elements:
                        self.processed_textbox_elements.append(id(tb_element))

                    _log.debug(
                        f"Found textbox content with {len(textbox_elements)} elements"
                    )
                    # 处理文本框内容并添加到已处理元素列表
                    tbc = self._handle_textbox_content(textbox_elements, doc)
                    added_elements.extend(tbc)

            # 检查表格元素
            if tag_name == "tbl":
                try:
                    # 处理表格元素
                    t = self._handle_tables(element, doc)
                    added_elements.extend(t)
                except Exception:
                    # 如果表格解析失败，记录调试信息
                    _log.debug("could not parse a table, broken docx table")
                    
            # 检查图像元素
            elif drawing_blip:
                # 处理图片元素
                pics = self._handle_pictures(drawing_blip, doc)
                added_elements.extend(pics)
                
                # 检查图像后的文本内容
                if (
                    tag_name == "p"
                    and element.find(
                        ".//w:t", namespaces=MsWordDocumentBackend._BLIP_NAMESPACES
                    )
                    is not None
                ):
                    # 处理文本元素
                    te1 = self._handle_text_elements(element, doc)
                    added_elements.extend(te1)
                    
            # 检查DrawingML元素
            elif drawingml_els:
                # 如果DOCX到PDF转换器尚未初始化，则获取转换器
                if (
                    self.docx_to_pdf_converter is None
                    and self.docx_to_pdf_converter_init is False
                ):
                    self.docx_to_pdf_converter = get_docx_to_pdf_converter()
                    self.docx_to_pdf_converter_init = True

                # 如果没有可用的转换器，则显示警告信息
                if self.docx_to_pdf_converter is None:
                    if self.display_drawingml_warning:
                        if self.docx_to_pdf_converter is None:
                            _log.warning(
                                "Found DrawingML elements in document, but no DOCX to PDF converters. "
                                "If you want these exported, make sure you have "
                                "LibreOffice binary in PATH or specify its path with DOCLING_LIBREOFFICE_CMD."
                            )
                            self.display_drawingml_warning = False
                else:
                    # 处理DrawingML元素
                    self._handle_drawingml(doc=doc, drawingml_els=drawingml_els)
                    
            # 检查SDT容器元素，如目录
            elif tag_name == "sdt":
                # 查找SDT内容元素
                sdt_content = element.find(
                    ".//w:sdtContent", namespaces=MsWordDocumentBackend._BLIP_NAMESPACES
                )
                if sdt_content is not None:
                    # 遍历<w:sdtContent>内的段落、运行或文本
                    paragraphs = sdt_content.findall(
                        ".//w:p", namespaces=MsWordDocumentBackend._BLIP_NAMESPACES
                    )
                    # 处理每个段落
                    for p in paragraphs:
                        te = self._handle_text_elements(p, doc)
                        added_elements.extend(te)
                        
            # 检查文本段落元素
            elif tag_name == "p":
                # 处理文本元素（包括段落属性如"tcPr", "sectPr"等）
                te = self._handle_text_elements(element, doc)
                added_elements.extend(te)
                
            # 忽略其他未知元素并记录日志
            else:
                _log.debug(f"Ignoring element in DOCX with tag: {tag_name}")

        # 返回更新后的文档和添加的元素引用列表
        return doc, added_elements

    def _str_to_int(
        self, s: Optional[str], default: Optional[int] = 0
    ) -> Optional[int]:
        """
        将字符串转换为整数。
        
        Args:
            s: 要转换的字符串
            default: 默认值，转换失败时返回
            
        Returns:
            Optional[int]: 转换后的整数，转换失败时返回默认值
        """
        if s is None:
            return None
        try:
            return int(s)
        except ValueError:
            return default

    def _split_text_and_number(self, input_string: str) -> list[str]:
        """
        分割字符串中的文本和数字部分。
        
        Args:
            input_string: 输入字符串
            
        Returns:
            list[str]: 分割后的部分列表
        """
        match = re.match(r"(\D+)(\d+)$|^(\d+)(\D+)", input_string)
        if match:
            parts = list(filter(None, match.groups()))
            return parts
        else:
            return [input_string]

    def _get_numId_and_ilvl(
        self, paragraph: Paragraph
    ) -> tuple[Optional[int], Optional[int]]:
        """
        获取段落的列表编号ID和层级。
        
        Args:
            paragraph: 段落对象
            
        Returns:
            tuple[Optional[int], Optional[int]]: (numId, ilvl) 元组
        """
        # 访问段落的XML元素
        numPr = paragraph._element.find(
            ".//w:numPr", namespaces=paragraph._element.nsmap
        )

        if numPr is not None:
            # 获取 numId 元素并提取值
            numId_elem = numPr.find("w:numId", namespaces=paragraph._element.nsmap)
            ilvl_elem = numPr.find("w:ilvl", namespaces=paragraph._element.nsmap)
            numId = numId_elem.get(self.XML_KEY) if numId_elem is not None else None
            ilvl = ilvl_elem.get(self.XML_KEY) if ilvl_elem is not None else None

            return self._str_to_int(numId, None), self._str_to_int(ilvl, None)

        return None, None  # 如果段落不是列表的一部分

    def _get_list_counter(self, numid: int, ilvl: int) -> int:
        """
        获取并递增特定 numId 和 ilvl 组合的计数器。
        
        Args:
            numid: 列表编号ID
            ilvl: 列表层级
            
        Returns:
            int: 当前计数器值
        """
        key = (numid, ilvl)
        if key not in self.list_counters:
            self.list_counters[key] = 0
        self.list_counters[key] += 1
        return self.list_counters[key]

    def _reset_list_counters_for_new_sequence(self, numid: int):
        """
        开始新的编号序列时重置计数器。
        
        Args:
            numid: 列表编号ID
        """
        # 重置此 numid 的所有计数器
        keys_to_reset = [key for key in self.list_counters.keys() if key[0] == numid]
        for key in keys_to_reset:
            self.list_counters[key] = 0

    def _is_numbered_list(self, numId: int, ilvl: int) -> bool:
        """
        根据 numFmt 值检查列表是否为编号列表。
        
        Args:
            numId: 列表编号ID
            ilvl: 列表层级
            
        Returns:
            bool: 如果是编号列表返回 True，否则返回 False
        """
        try:
            # 访问文档的编号部分
            if not hasattr(self.docx_obj, "part") or not hasattr(
                self.docx_obj.part, "package"
            ):
                return False

            numbering_part = None
            # 查找编号部分
            for part in self.docx_obj.part.package.parts:
                if "numbering" in part.partname:
                    numbering_part = part
                    break

            if numbering_part is None:
                return False

            # 解析编号 XML
            numbering_root = numbering_part.element
            namespaces = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            }

            # 查找具有给定 numId 的编号定义
            num_xpath = f".//w:num[@w:numId='{numId}']"
            num_element = numbering_root.find(num_xpath, namespaces=namespaces)

            if num_element is None:
                return False

            # 从 num 元素获取 abstractNumId
            abstract_num_id_elem = num_element.find(
                ".//w:abstractNumId", namespaces=namespaces
            )
            if abstract_num_id_elem is None:
                return False

            abstract_num_id = abstract_num_id_elem.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
            )
            if abstract_num_id is None:
                return False

            # 查找抽象编号定义
            abstract_num_xpath = (
                f".//w:abstractNum[@w:abstractNumId='{abstract_num_id}']"
            )
            abstract_num_element = numbering_root.find(
                abstract_num_xpath, namespaces=namespaces
            )

            if abstract_num_element is None:
                return False

            # 查找给定 ilvl 的层级定义
            lvl_xpath = f".//w:lvl[@w:ilvl='{ilvl}']"
            lvl_element = abstract_num_element.find(lvl_xpath, namespaces=namespaces)

            if lvl_element is None:
                return False

            # 获取 numFmt 元素
            num_fmt_element = lvl_element.find(".//w:numFmt", namespaces=namespaces)
            if num_fmt_element is None:
                return False

            num_fmt = num_fmt_element.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
            )

            # 编号格式包括: decimal, lowerRoman, upperRoman, lowerLetter, upperLetter
            # 项目符号格式包括: bullet
            numbered_formats = {
                "decimal",
                "lowerRoman",
                "upperRoman",
                "lowerLetter",
                "upperLetter",
                "decimalZero",
            }

            return num_fmt in numbered_formats

        except Exception as e:
            _log.debug(f"Error determining if list is numbered: {e}")
            return False

    def _get_heading_and_level(self, style_label: str) -> tuple[str, Optional[int]]:
        """
        从样式标签获取标题和层级。
        
        Args:
            style_label: 样式标签
            
        Returns:
            tuple[str, Optional[int]]: (标签字符串, 层级) 元组
        """
        parts = self._split_text_and_number(style_label)

        if len(parts) == 2:
            parts.sort()
            label_str: str = ""
            label_level: Optional[int] = 0
            if parts[0].strip().lower() == "heading":
                label_str = "Heading"
                label_level = self._str_to_int(parts[1], None)
            if parts[1].strip().lower() == "heading":
                label_str = "Heading"
                label_level = self._str_to_int(parts[0], None)
            return label_str, label_level

        return style_label, None

    def _get_label_and_level(self, paragraph: Paragraph) -> tuple[str, Optional[int]]:
        """
        获取段落的标签和层级。
        
        Args:
            paragraph: 段落对象
            
        Returns:
            tuple[str, Optional[int]]: (标签, 层级) 元组
        """
        if paragraph.style is None:
            return "Normal", None

        label = paragraph.style.style_id
        name = paragraph.style.name
        base_style_label = None
        base_style_name = None
        if base_style := getattr(paragraph.style, "base_style", None):
            base_style_label = base_style.style_id
            base_style_name = base_style.name

        if label is None:
            return "Normal", None

        if ":" in label:
            parts = label.split(":")
            if len(parts) == 2:
                return parts[0], self._str_to_int(parts[1], None)

        if "heading" in label.lower():
            return self._get_heading_and_level(label)
        if "heading" in name.lower():
            return self._get_heading_and_level(name)
        if base_style_label and "heading" in base_style_label.lower():
            return self._get_heading_and_level(base_style_label)
        if base_style_name and "heading" in base_style_name.lower():
            return self._get_heading_and_level(base_style_name)

        return label, None

    @classmethod
    def _get_format_from_run(cls, run: Run) -> Optional[Formatting]:
        """
        从 Run 对象获取格式信息。
        
        Args:
            run: Run 对象
            
        Returns:
            Optional[Formatting]: 格式对象
        """
        # .bold 和 .italic 属性是布尔值，但 .underline 可能是枚举
        # 如 WD_UNDERLINE.THICK (值为 6)，所以需要转换为布尔值
        is_bold = run.bold or False
        is_italic = run.italic or False
        is_strikethrough = run.font.strike or False
        # 将任何非 None 的下划线值转换为 True
        is_underline = bool(run.underline is not None and run.underline)
        is_sub = run.font.subscript or False
        is_sup = run.font.superscript or False
        script = Script.SUB if is_sub else Script.SUPER if is_sup else Script.BASELINE

        return Formatting(
            bold=is_bold,
            italic=is_italic,
            underline=is_underline,
            strikethrough=is_strikethrough,
            script=script,
        )

    def _get_paragraph_elements(self, paragraph: Paragraph):
        """
        提取段落元素及其格式和超链接信息。
        
        Args:
            paragraph: 段落对象
            
        Returns:
            list[tuple[str, Optional[Formatting], Optional[Union[AnyUrl, Path]]]]: 
            段落元素列表，每个元素包含文本、格式和超链接信息
        """

        # 目前保留空段落以保持向后兼容性:
        if paragraph.text.strip() == "":
            return [("", None, None)]

        paragraph_elements: list[
            tuple[str, Optional[Formatting], Optional[Union[AnyUrl, Path]]]
        ] = []
        group_text = ""
        previous_format = None

        # 遍历段落的 runs 并按格式分组
        for c in paragraph.iter_inner_content():
            if isinstance(c, Hyperlink):
                text = c.text
                hyperlink = Path(c.address)
                format = (
                    self._get_format_from_run(c.runs[0])
                    if c.runs and len(c.runs) > 0
                    else None
                )
            elif isinstance(c, Run):
                text = c.text
                hyperlink = None
                format = self._get_format_from_run(c)
            else:
                continue

            if (len(text.strip()) and format != previous_format) or (
                hyperlink is not None
            ):
                # 如果非空文本的样式发生变化，则添加前一个组
                if len(group_text.strip()) > 0:
                    paragraph_elements.append(
                        (group_text.strip(), previous_format, None)
                    )
                group_text = ""

                # 如果有超链接，则立即添加
                if hyperlink is not None:
                    paragraph_elements.append((text.strip(), format, hyperlink))
                    text = ""
                else:
                    previous_format = format

            group_text += text

        # 格式化最后一个组
        if len(group_text.strip()) > 0:
            paragraph_elements.append((group_text.strip(), format, None))

        return paragraph_elements

    def _get_paragraph_position(self, paragraph_element):
        """
        从段落元素中提取垂直位置信息。
        
        Args:
            paragraph_element: 段落元素
            
        Returns:
            段落的垂直位置信息
        """
        # 首先尝试直接从具有顺序相关属性的 w:p 元素获取索引
        if (
            hasattr(paragraph_element, "getparent")
            and paragraph_element.getparent() is not None
        ):
            parent = paragraph_element.getparent()
            # 获取所有段落兄弟元素
            paragraphs = [
                p for p in parent.getchildren() if etree.QName(p).localname == "p"
            ]
            # 查找当前段落在其兄弟元素中的索引
            try:
                paragraph_index = paragraphs.index(paragraph_element)
                return paragraph_index  # 使用索引作为位置以保持一致的排序
            except ValueError:
                pass

        # 在元素属性和祖先元素中查找位置提示
        for elem in (*[paragraph_element], *paragraph_element.iterancestors()):
            # 检查直接位置属性
            for attr_name in ["y", "top", "positionY", "y-position", "position"]:
                value = elem.get(attr_name)
                if value:
                    try:
                        # 移除任何非数字字符（如 'pt', 'px' 等）
                        clean_value = re.sub(r"[^0-9.]", "", value)
                        if clean_value:
                            return float(clean_value)
                    except (ValueError, TypeError):
                        pass

            # 检查 transform 属性中的位置
            transform = elem.get("transform")
            if transform:
                # 从变换矩阵中提取平移分量
                match = re.search(r"translate\([^,]+,\s*([0-9.]+)", transform)
                if match:
                    try:
                        return float(match.group(1))
                    except ValueError:
                        pass

            # 检查 Word 格式中的锚点或相对位置指示器
            # 'dist' 属性可以指示相对定位
            for attr_name in ["distT", "distB", "anchor", "relativeFrom"]:
                if elem.get(attr_name) is not None:
                    return elem.sourceline  # 使用 XML 源代码行号作为后备

        # 对于 VML 形状，查找特定属性
        for ns_uri in paragraph_element.nsmap.values():
            if "vml" in ns_uri:
                # 尝试从 style 属性中提取位置
                style = paragraph_element.get("style")
                if style:
                    match = re.search(r"top:([0-9.]+)pt", style)
                    if match:
                        try:
                            return float(match.group(1))
                        except ValueError:
                            pass

        # 如果没有找到更好的位置指示器，则使用 XML 源代码行号作为顺序的代理
        return (
            paragraph_element.sourceline
            if hasattr(paragraph_element, "sourceline")
            else None
        )

    def _collect_textbox_paragraphs(self, textbox_elements):
        """
        从文本框元素中收集和组织段落。
        
        Args:
            textbox_elements: 文本框元素列表
            
        Returns:
            dict: 容器段落字典
        """
        processed_paragraphs = []
        container_paragraphs = {}

        for element in textbox_elements:
            element_id = id(element)
            # 如果已经处理过这个确切元素则跳过
            if element_id in processed_paragraphs:
                continue

            tag_name = etree.QName(element).localname
            processed_paragraphs.append(element_id)

            # 处理直接找到的段落 (VML 文本框)
            if tag_name == "p":
                # 查找包含的文本框或形状元素
                container_id = None
                for ancestor in element.iterancestors():
                    if any(ns in ancestor.tag for ns in ["textbox", "shape", "txbx"]):
                        container_id = id(ancestor)
                        break

                if container_id not in container_paragraphs:
                    container_paragraphs[container_id] = []
                container_paragraphs[container_id].append(
                    (element, self._get_paragraph_position(element))
                )

            # 处理 txbxContent 元素 (Word DrawingML 文本框)
            elif tag_name == "txbxContent":
                paragraphs = element.findall(".//w:p", namespaces=element.nsmap)
                container_id = id(element)
                if container_id not in container_paragraphs:
                    container_paragraphs[container_id] = []

                for p in paragraphs:
                    p_id = id(p)
                    if p_id not in processed_paragraphs:
                        processed_paragraphs.append(p_id)
                        container_paragraphs[container_id].append(
                            (p, self._get_paragraph_position(p))
                        )
            else:
                # 尝试从未知元素中提取任何段落
                paragraphs = element.findall(".//w:p", namespaces=element.nsmap)
                container_id = id(element)
                if container_id not in container_paragraphs:
                    container_paragraphs[container_id] = []

                for p in paragraphs:
                    p_id = id(p)
                    if p_id not in processed_paragraphs:
                        processed_paragraphs.append(p_id)
                        container_paragraphs[container_id].append(
                            (p, self._get_paragraph_position(p))
                        )

        return container_paragraphs

    def _handle_textbox_content(
        self,
        textbox_elements: list,
        doc: DoclingDocument,
    ) -> list[RefItem]:
        """
        处理文本框内容并将其添加到文档结构中。
        
        Args:
            textbox_elements: 文本框元素列表
            doc: DoclingDocument 对象
            
        Returns:
            list[RefItem]: 元素引用列表
        """
        elem_ref: list[RefItem] = []
        level = self._get_level()
        # 创建一个文本框组来包含文本框中的所有文本
        textbox_group = doc.add_group(
            label=GroupLabel.SECTION,
            parent=self.parents[level - 1],
            name="textbox",
            content_layer=self.content_layer,
        )
        elem_ref.append(textbox_group.get_ref())
        # 将此设置为当前父级以确保文本框内容
        # 在文档结构中正确嵌套
        original_parent = self.parents[level]
        self.parents[level] = textbox_group

        # 收集和组织段落
        container_paragraphs = self._collect_textbox_paragraphs(textbox_elements)

        # 处理所有段落
        all_paragraphs = []

        # 在每个容器内排序段落，然后处理容器
        for paragraphs in container_paragraphs.values():
            # 按垂直位置在每个容器内排序
            sorted_container_paragraphs = sorted(
                paragraphs,
                key=lambda x: (
                    x[1] is None,
                    x[1] if x[1] is not None else float("inf"),
                ),
            )

            # 将排序后的段落添加到我们的处理列表中
            all_paragraphs.extend(sorted_container_paragraphs)

        # 跟踪已处理的段落以避免重复（相同内容和位置）
        processed_paragraphs = set()

        # 处理所有段落
        for p, position in all_paragraphs:
            # 创建段落对象以获取文本内容
            paragraph = Paragraph(p, self.docx_obj)
            text_content = paragraph.text

            # 基于内容和位置创建唯一标识符
            paragraph_id = (text_content, position)

            # 如果此段落（相同内容和位置）已被处理则跳过
            if paragraph_id in processed_paragraphs:
                _log.debug(
                    f"Skipping duplicate paragraph: content='{text_content[:50]}...', position={position}"
                )
                continue

            # 标记此段落已处理
            processed_paragraphs.add(paragraph_id)

            elem_ref.extend(self._handle_text_elements(p, doc))

        # 恢复原始父级
        self.parents[level] = original_parent
        return elem_ref

    def _handle_equations_in_text(self, element, text):
        """
        处理文本中的公式。
        
        Args:
            element: 元素对象
            text: 文本内容
            
        Returns:
            tuple: (处理后的文本, 公式列表)
        """
        only_texts = []
        only_equations = []
        texts_and_equations = []
        for subt in element.iter():
            tag_name = etree.QName(subt).localname
            if tag_name == "t" and "math" not in subt.tag:
                if isinstance(subt.text, str):
                    only_texts.append(subt.text)
                    texts_and_equations.append(subt.text)
            elif "oMath" in subt.tag and "oMathPara" not in subt.tag:
                latex_equation = str(oMath2Latex(subt)).strip()
                if len(latex_equation) > 0:
                    only_equations.append(
                        self.equation_bookends.format(EQ=latex_equation)
                    )
                    texts_and_equations.append(
                        self.equation_bookends.format(EQ=latex_equation)
                    )

        if len(only_equations) < 1:
            return text, []

        if (
            re.sub(r"\s+", "", "".join(only_texts)).strip()
            != re.sub(r"\s+", "", text).strip()
        ):
            # 如果我们无法重构初始原始文本
            # 不要尝试解析公式并返回原始文本
            return text, []

        # 将公式插入原始文本中
        # 这样做是为了保持空白结构
        output_text = text[:]
        init_i = 0
        for i_substr, substr in enumerate(texts_and_equations):
            if len(substr) == 0:
                continue

            if substr in output_text[init_i:]:
                init_i += output_text[init_i:].find(substr) + len(substr)
            else:
                if i_substr > 0:
                    output_text = output_text[:init_i] + substr + output_text[init_i:]
                    init_i += len(substr)
                else:
                    output_text = substr + output_text

        return output_text, only_equations

    def _create_or_reuse_parent(
        self,
        *,
        doc: DoclingDocument,
        prev_parent: Optional[NodeItem],
        paragraph_elements: list,
    ) -> Optional[NodeItem]:
        """
        为段落元素创建或重用父级节点。
        
        Args:
            doc: DoclingDocument 对象
            prev_parent: 前一个父级节点
            paragraph_elements: 段落元素列表
            
        Returns:
            Optional[NodeItem]: 父级节点
        """
        return (
            doc.add_inline_group(parent=prev_parent, content_layer=self.content_layer)
            if len(paragraph_elements) > 1
            else prev_parent
        )

    def _handle_text_elements(
        self,
        element: BaseOxmlElement,
        doc: DoclingDocument,
    ) -> list[RefItem]:
        """
        处理文本元素。
        
        Args:
            element: 元素对象
            doc: DoclingDocument 对象
            
        Returns:
            list[RefItem]: 元素引用列表
        """
        elem_ref: list[RefItem] = []
        paragraph = Paragraph(element, self.docx_obj)
        paragraph_elements = self._get_paragraph_elements(paragraph)
        text, equations = self._handle_equations_in_text(
            element=element, text=paragraph.text
        )

        if text is None:
            return elem_ref
        text = text.strip()

        # 常见的项目符号和编号列表样式。
        # "List Bullet", "List Number", "List Paragraph"
        # 识别列表是否为编号列表
        p_style_id, p_level = self._get_label_and_level(paragraph)
        numid, ilevel = self._get_numId_and_ilvl(paragraph)

        if numid == 0:
            numid = None

        # 处理列表
        if (
            numid is not None
            and ilevel is not None
            and p_style_id not in ["Title", "Heading"]
        ):
            # 通过检查 numFmt 来确认这是否实际上是编号列表
            is_numbered = self._is_numbered_list(numid, ilevel)

            li = self._add_list_item(
                doc=doc,
                numid=numid,
                ilevel=ilevel,
                elements=paragraph_elements,
                is_numbered=is_numbered,
            )
            elem_ref.extend(li)  # 必须是引用!!!
            self._update_history(p_style_id, p_level, numid, ilevel)
            return elem_ref
        elif (
            numid is None
            and self._prev_numid() is not None
            and p_style_id not in ["Title", "Heading"]
        ):  # 关闭列表
            if self.level_at_new_list:
                for key in range(len(self.parents)):
                    if key >= self.level_at_new_list:
                        self.parents[key] = None
                self.level = self.level_at_new_list - 1
                self.level_at_new_list = None
            else:
                for key in range(len(self.parents)):
                    self.parents[key] = None
                self.level = 0

        if p_style_id in ["Title"]:
            for key in range(len(self.parents)):
                self.parents[key] = None
            te = doc.add_text(
                parent=None,
                label=DocItemLabel.TITLE,
                text=text,
                content_layer=self.content_layer,
            )
            self.parents[0] = te
            elem_ref.append(te.get_ref())
        elif "Heading" in p_style_id:
            style_element = getattr(paragraph.style, "element", None)
            if style_element is not None:
                is_numbered_style = (
                    "<w:numPr>" in style_element.xml or "<w:numPr>" in element.xml
                )
            else:
                is_numbered_style = False
            h1 = self._add_heading(doc, p_level, text, is_numbered_style)
            elem_ref.extend(h1)

        elif len(equations) > 0:
            if (paragraph.text is None or len(paragraph.text.strip()) == 0) and len(
                text
            ) > 0:
                # 独立公式
                level = self._get_level()
                t1 = doc.add_text(
                    label=DocItemLabel.FORMULA,
                    parent=self.parents[level - 1],
                    text=text.replace("<eq>", "").replace("</eq>", ""),
                    content_layer=self.content_layer,
                )
                elem_ref.append(t1.get_ref())
            else:
                # 行内公式
                level = self._get_level()
                inline_equation = doc.add_inline_group(
                    parent=self.parents[level - 1], content_layer=self.content_layer
                )
                elem_ref.append(inline_equation.get_ref())
                text_tmp = text
                for eq in equations:
                    if len(text_tmp) == 0:
                        break

                    split_text_tmp = text_tmp.split(eq.strip(), maxsplit=1)

                    pre_eq_text = split_text_tmp[0]
                    text_tmp = "" if len(split_text_tmp) == 1 else split_text_tmp[1]

                    if len(pre_eq_text) > 0:
                        e1 = doc.add_text(
                            label=DocItemLabel.TEXT,
                            parent=inline_equation,
                            text=pre_eq_text,
                            content_layer=self.content_layer,
                        )
                        elem_ref.append(e1.get_ref())
                    e2 = doc.add_text(
                        label=DocItemLabel.FORMULA,
                        parent=inline_equation,
                        text=eq.replace("<eq>", "").replace("</eq>", ""),
                        content_layer=self.content_layer,
                    )
                    elem_ref.append(e2.get_ref())

                if len(text_tmp) > 0:
                    e3 = doc.add_text(
                        label=DocItemLabel.TEXT,
                        parent=inline_equation,
                        text=text_tmp.strip(),
                        content_layer=self.content_layer,
                    )
                    elem_ref.append(e3.get_ref())

        elif p_style_id in [
            "Paragraph",
            "Normal",
            "Subtitle",
            "Author",
            "DefaultText",
            "ListParagraph",
            "ListBullet",
            "Quote",
        ]:
            level = self._get_level()
            parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents.get(level - 1),
                paragraph_elements=paragraph_elements,
            )
            for text, format, hyperlink in paragraph_elements:
                t2 = doc.add_text(
                    label=DocItemLabel.TEXT,
                    parent=parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                    content_layer=self.content_layer,
                )
                elem_ref.append(t2.get_ref())

        else:
            # 文本样式名称不仅有默认值，还可能有用户自定义值
            # 因此我们将所有其他标签视为纯文本
            level = self._get_level()
            parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents.get(level - 1),
                paragraph_elements=paragraph_elements,
            )
            for text, format, hyperlink in paragraph_elements:
                t3 = doc.add_text(
                    label=DocItemLabel.TEXT,
                    parent=parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                    content_layer=self.content_layer,
                )
                elem_ref.append(t3.get_ref())

        self._update_history(p_style_id, p_level, numid, ilevel)
        return elem_ref

    def _add_heading(
        self,
        doc: DoclingDocument,
        curr_level: Optional[int],
        text: str,
        is_numbered_style: bool = False,
    ) -> list[RefItem]:
        """
        添加标题。
        
        Args:
            doc: DoclingDocument 对象
            curr_level: 当前层级
            text: 标题文本
            is_numbered_style: 是否为编号样式
            
        Returns:
            list[RefItem]: 元素引用列表
        """
        elem_ref: list[RefItem] = []
        level = self._get_level()
        if isinstance(curr_level, int):
            if curr_level > level:
                # 添加不可见组
                for i in range(level, curr_level):
                    gr1 = doc.add_group(
                        parent=self.parents[i - 1],
                        label=GroupLabel.SECTION,
                        name=f"header-{i}",
                    )
                    elem_ref.append(gr1.get_ref())
                    self.parents[i] = gr1

            elif curr_level < level:
                # 移除尾部
                for key in range(len(self.parents)):
                    if key >= curr_level:
                        self.parents[key] = None

            current_level = curr_level
            parent_level = curr_level - 1
            add_level = curr_level
        else:
            current_level = self.level
            parent_level = self.level - 1
            add_level = 1

        if is_numbered_style:
            if add_level in self.numbered_headers:
                self.numbered_headers[add_level] += 1
            else:
                self.numbered_headers[add_level] = 1
            text = f"{self.numbered_headers[add_level]} {text}"

            # 重置更深层级
            next_level = add_level + 1
            while next_level in self.numbered_headers:
                self.numbered_headers[next_level] = 0
                next_level += 1

            # 扫描上层级别
            previous_level = add_level - 1
            while previous_level in self.numbered_headers:
                # MSWord 约定: 不允许空的子层级
                # 即，没有子章节(2.1)的子子章节(2.0.1)
                # 被处理为 2.1.1
                if self.numbered_headers[previous_level] == 0:
                    self.numbered_headers[previous_level] += 1

                text = f"{self.numbered_headers[previous_level]}.{text}"
                previous_level -= 1

        hd = doc.add_heading(
            parent=self.parents[parent_level],
            text=text,
            level=add_level,
        )
        self.parents[current_level] = hd
        elem_ref.append(hd.get_ref())
        return elem_ref

    def _add_formatted_list_item(
        self,
        doc: DoclingDocument,
        elements: list,
        marker: str,
        enumerated: bool,
        level: int,
    ) -> list[RefItem]:
        """
        添加格式化的列表项。
        
        Args:
            doc: DoclingDocument 对象
            elements: 元素列表
            marker: 标记
            enumerated: 是否编号
            level: 层级
            
        Returns:
            list[RefItem]: 元素引用列表
        """
        elem_ref: list[RefItem] = []
        # 这在构造上不应该发生
        if not isinstance(self.parents[level], ListGroup):
            return elem_ref
        if not elements:
            return elem_ref

        if len(elements) == 1:
            text, format, hyperlink = elements[0]
            if text:
                doc.add_list_item(
                    marker=marker,
                    enumerated=enumerated,
                    parent=self.parents[level],
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                )
        else:
            new_item = doc.add_list_item(
                marker=marker,
                enumerated=enumerated,
                parent=self.parents[level],
                text="",
            )
            new_parent = doc.add_inline_group(parent=new_item)
            for text, format, hyperlink in elements:
                if text:
                    doc.add_text(
                        label=DocItemLabel.TEXT,
                        parent=new_parent,
                        text=text,
                        formatting=format,
                        hyperlink=hyperlink,
                        content_layer=self.content_layer,
                    )
        return elem_ref

    def _add_list_item(
        self,
        *,
        doc: DoclingDocument,
        numid: int,
        ilevel: int,
        elements: list,
        is_numbered: bool = False,
    ) -> list[RefItem]:
        """
        添加列表项。
        
        Args:
            doc: DoclingDocument 对象
            numid: 列表编号ID
            ilevel: 列表层级
            elements: 元素列表
            is_numbered: 是否编号
            
        Returns:
            list[RefItem]: 元素引用列表
        """
        elem_ref: list[RefItem] = []
        # 此方法始终使用 is_numbered 调用。应正确处理编号列表。
        if not elements:
            return elem_ref
        enum_marker = ""

        level = self._get_level()
        prev_indent = self._prev_indent()
        if self._prev_numid() is None:  # 打开新列表
            self.level_at_new_list = level

            # 为新编号序列重置计数器
            self._reset_list_counters_for_new_sequence(numid)

            list_gr = doc.add_list_group(
                name="list",
                parent=self.parents[level - 1],
                content_layer=self.content_layer,
            )
            self.parents[level] = list_gr
            elem_ref.append(list_gr.get_ref())

            # 如果这是枚举元素，则设置标记和枚举参数。
            if is_numbered:
                counter = self._get_list_counter(numid, ilevel)
                enum_marker = str(counter) + "."
            else:
                enum_marker = ""
            self._add_formatted_list_item(
                doc, elements, enum_marker, is_numbered, level
            )
        elif (
            self._prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and prev_indent < ilevel
        ):  # 打开缩进列表
            for i in range(
                self.level_at_new_list + prev_indent + 1,
                self.level_at_new_list + ilevel + 1,
            ):
                list_gr1 = doc.add_list_group(
                    name="list",
                    parent=self.parents[i - 1],
                    content_layer=self.content_layer,
                )
                self.parents[i] = list_gr1
                elem_ref.append(list_gr1.get_ref())

            # TODO: 如果这是枚举元素，则设置标记和枚举参数。
            if is_numbered:
                counter = self._get_list_counter(numid, ilevel)
                enum_marker = str(counter) + "."
            else:
                enum_marker = ""
            self._add_formatted_list_item(
                doc,
                elements,
                enum_marker,
                is_numbered,
                self.level_at_new_list + ilevel,
            )
        elif (
            self._prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and ilevel < prev_indent
        ):  # 关闭列表
            for k in self.parents:
                if k > self.level_at_new_list + ilevel:
                    self.parents[k] = None

            # TODO: 如果这是枚举元素，则设置标记和枚举参数。
            if is_numbered:
                counter = self._get_list_counter(numid, ilevel)
                enum_marker = str(counter) + "."
            else:
                enum_marker = ""
            self._add_formatted_list_item(
                doc,
                elements,
                enum_marker,
                is_numbered,
                self.level_at_new_list + ilevel,
            )

        elif self._prev_numid() == numid or prev_indent == ilevel:
            # 如果这是枚举元素，则设置标记和枚举参数。
            if is_numbered:
                counter = self._get_list_counter(numid, ilevel)
                enum_marker = str(counter) + "."
            else:
                enum_marker = ""
            self._add_formatted_list_item(
                doc, elements, enum_marker, is_numbered, level - 1
            )
        return elem_ref

    @staticmethod
    def _group_cell_elements(
        group_name: str,
        doc: DoclingDocument,
        provs_in_cell: list[RefItem],
        docling_table: TableItem,
        content_layer: ContentLayer = ContentLayer.BODY,
    ) -> RefItem:
        """
        将单元格中的元素分组。
        
        Args:
            group_name: 组名称
            doc: DoclingDocument 对象
            provs_in_cell: 单元格中的元素引用列表
            docling_table: 表格项
            content_layer: 内容层
            
        Returns:
            RefItem: 组的引用
        """
        group_element = doc.add_group(
            label=GroupLabel.UNSPECIFIED,
            name=group_name,
            parent=docling_table,
            content_layer=content_layer,
        )
        for prov in provs_in_cell:
            group_element.children.append(prov)
            pr_item = prov.resolve(doc)
            item_parent = pr_item.parent.resolve(doc)
            if pr_item.get_ref() in item_parent.children:
                item_parent.children.remove(pr_item.get_ref())
            pr_item.parent = group_element.get_ref()
        ref_for_rich_cell = group_element.get_ref()
        return ref_for_rich_cell

    def _handle_tables(
        self,
        element: BaseOxmlElement,
        doc: DoclingDocument,
    ) -> list[RefItem]:
        """
        处理表格。
        
        Args:
            element: 元素对象
            doc: DoclingDocument 对象
            
        Returns:
            list[RefItem]: 元素引用列表
        """
        elem_ref: list[RefItem] = []
        table: Table = Table(element, self.docx_obj)
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        _log.debug(f"Table grid with {num_rows} rows and {num_cols} columns")

        if num_rows == 1 and num_cols == 1:
            cell_element = table.rows[0].cells[0]
            # 如果我们有一个只有1个单元格的表格，我们认为它是装饰元素
            # 并继续将单元格的内容作为文档正文处理
            self._walk_linear(cell_element._element, doc)
            return elem_ref

        data = TableData(num_rows=num_rows, num_cols=num_cols)
        level = self._get_level()
        docling_table = doc.add_table(
            data=data, parent=self.parents[level - 1], content_layer=self.content_layer
        )
        elem_ref.append(docling_table.get_ref())

        cell_set: set[CT_Tc] = set()
        for row_idx, row in enumerate(table.rows):
            _log.debug(f"Row index {row_idx} with {len(row.cells)} populated cells")
            col_idx = 0
            while col_idx < num_cols:
                cell: _Cell = row.cells[col_idx]
                _log.debug(
                    f" col {col_idx} grid_span {cell.grid_span} grid_cols_before {row.grid_cols_before}"
                )
                if cell is None or cell._tc in cell_set:
                    _log.debug("  skipped since repeated content")
                    col_idx += cell.grid_span
                    continue
                else:
                    cell_set.add(cell._tc)

                spanned_idx = row_idx
                spanned_tc: Optional[CT_Tc] = cell._tc
                while spanned_tc == cell._tc:
                    spanned_idx += 1
                    spanned_tc = (
                        table.rows[spanned_idx].cells[col_idx]._tc
                        if spanned_idx < num_rows
                        else None
                    )
                _log.debug(f"  spanned before row {spanned_idx}")

                # 检测单元格文本中的公式
                text, equations = self._handle_equations_in_text(
                    element=cell._element, text=cell.text
                )
                if len(equations) == 0:
                    text = cell.text
                else:
                    text = text.replace("<eq>", "$").replace("</eq>", "$")

                provs_in_cell: list[RefItem] = []
                rich_table_cell: bool = self._is_rich_table_cell(cell)

                if rich_table_cell:
                    _, provs_in_cell = self._walk_linear(cell._element, doc)
                _log.debug(f"Table cell {row_idx},{col_idx} rich? {rich_table_cell}")

                if len(provs_in_cell) > 0:
                    # 单元格有多个元素，我们需要将它们分组
                    rich_table_cell = True
                    group_name = f"rich_cell_group_{len(doc.tables)}_{col_idx}_{row.grid_cols_before + row_idx}"
                    ref_for_rich_cell = MsWordDocumentBackend._group_cell_elements(
                        group_name,
                        doc,
                        provs_in_cell,
                        docling_table,
                        content_layer=self.content_layer,
                    )

                if rich_table_cell:
                    rich_cell = RichTableCell(
                        text=text,
                        row_span=spanned_idx - row_idx,
                        col_span=cell.grid_span,
                        start_row_offset_idx=row.grid_cols_before + row_idx,
                        end_row_offset_idx=row.grid_cols_before + spanned_idx,
                        start_col_offset_idx=col_idx,
                        end_col_offset_idx=col_idx + cell.grid_span,
                        column_header=row.grid_cols_before + row_idx == 0,
                        row_header=False,
                        ref=ref_for_rich_cell,  # 指向围绕子元素的人工组
                    )
                    doc.add_table_cell(table_item=docling_table, cell=rich_cell)
                    col_idx += cell.grid_span
                else:
                    simple_cell = TableCell(
                        text=text,
                        row_span=spanned_idx - row_idx,
                        col_span=cell.grid_span,
                        start_row_offset_idx=row.grid_cols_before + row_idx,
                        end_row_offset_idx=row.grid_cols_before + spanned_idx,
                        start_col_offset_idx=col_idx,
                        end_col_offset_idx=col_idx + cell.grid_span,
                        column_header=row.grid_cols_before + row_idx == 0,
                        row_header=False,
                    )
                    doc.add_table_cell(table_item=docling_table, cell=simple_cell)
                    col_idx += cell.grid_span
        return elem_ref

    def _has_blip(self, element: BaseOxmlElement) -> bool:
        """
        检查 docx 元素是否包含任何 BLIP 作为子元素。

        Args:
            element: docx 元素

        Returns:
            bool: 元素是否包含 BLIP 作为直接子元素
        """

        for item in element:
            if self.blip_xpath_expr(item):
                return True
            if item.findall(
                ".//w:drawing", namespaces=MsWordDocumentBackend._BLIP_NAMESPACES
            ):
                return True

        return False

    def _is_rich_table_cell(self, cell: _Cell) -> bool:
        """
        确定 docx 单元格是否应解析为 Docling RichTableCell。

        docx 单元格可以包含丰富内容并使用 Docling RichTableCell 解析。
        但这需要遍历 lxml 元素并创建节点项。如果单元格只包含纯文本，
        则使用 TableCell 解析更简单，也更受推荐。

        纯文本意味着：
        - 单元格只有一个段落
        - 段落仅由没有运行属性的 runs 组成
          (不需要 Docling 格式化)。
        - 单元格元素内没有其他块级元素。

        Args:
            cell: docx 单元格

        Returns:
            bool: docx 单元格是否应解析为 RichTableCell
        """
        tc = cell._tc

        # 必须只包含一个段落
        paragraphs = list(
            tc.iterchildren(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
            )
        )
        if len(paragraphs) > 1:
            return True

        # 没有其他内容
        allowed_tags = {"p", "tcPr"}  # 段落或表格单元格属性
        for child in tc:
            tag = child.tag.split("}")[-1]
            if tag not in allowed_tags:
                return True
        if self._has_blip(tc):
            return True

        # 段落必须包含没有运行属性的 runs
        for para in paragraphs:
            runs = list(
                para.iterchildren(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
                )
            )
            for rn in runs:
                item: Run = Run(rn, self.docx_obj)
                if item is not None:
                    fm = MsWordDocumentBackend._get_format_from_run(item)
                    if fm != Formatting():
                        return True

        # 所有检查通过：仅纯文本
        return False

    def _handle_pictures(
        self, drawing_blip: Any, doc: DoclingDocument
    ) -> list[RefItem]:
        """
        处理图片。
        
        Args:
            drawing_blip: 绘图 blip 对象
            doc: DoclingDocument 对象
            
        Returns:
            list[RefItem]: 元素引用列表
        """
        def get_docx_image(drawing_blip: Any) -> Optional[bytes]:
            """
            获取 DOCX 图像数据。
            
            Args:
                drawing_blip: 绘图 blip 对象
                
            Returns:
                Optional[bytes]: 图像数据
            """
            image_data: Optional[bytes] = None
            rId = drawing_blip[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if rId in self.docx_obj.part.rels:
                # 使用关系 ID 访问图像部分
                image_part = self.docx_obj.part.rels[rId].target_part
                image_data = image_part.blob  # 获取二进制图像数据
            return image_data

        elem_ref: list[RefItem] = []
        level = self._get_level()
        # 使用 PIL 打开 BytesIO 对象创建图像
        image_data: Optional[bytes] = get_docx_image(drawing_blip)
        if image_data is None:
            _log.warning("Warning: image cannot be found")
            p1 = doc.add_picture(
                parent=self.parents[level - 1],
                caption=None,
                content_layer=self.content_layer,
            )
            elem_ref.append(p1.get_ref())
        else:
            try:
                image_bytes = BytesIO(image_data)
                pil_image = Image.open(image_bytes)
                p2 = doc.add_picture(
                    parent=self.parents[level - 1],
                    image=ImageRef.from_pil(image=pil_image, dpi=72),
                    caption=None,
                    content_layer=self.content_layer,
                )
                elem_ref.append(p2.get_ref())
            except (UnidentifiedImageError, OSError):
                _log.warning("Warning: image cannot be loaded by Pillow")
                p3 = doc.add_picture(
                    parent=self.parents[level - 1],
                    caption=None,
                    content_layer=self.content_layer,
                )
                elem_ref.append(p3.get_ref())
        return elem_ref

    def _handle_drawingml(self, doc: DoclingDocument, drawingml_els: Any):
        """
        处理 DrawingML 元素。
        
        Args:
            doc: DoclingDocument 对象
            drawingml_els: DrawingML 元素
        """
        # 1) 创建原始文档的空副本
        dml_doc = self.load_msword_file(self.path_or_stream, self.document_hash)
        body = dml_doc._element.body
        for child in list(body):
            body.remove(child)

        # 2) 将 DrawingML 添加到空文档中
        new_para = dml_doc.add_paragraph()
        new_r = new_para.add_run()
        for dml in drawingml_els:
            new_r._r.append(deepcopy(dml))

        # 3) 导出 DOCX->PDF->PNG 并保存到 DoclingDocument 中
        level = self._get_level()
        try:
            pil_image = get_pil_from_dml_docx(
                dml_doc, converter=self.docx_to_pdf_converter
            )
            if pil_image is None:
                raise UnidentifiedImageError

            doc.add_picture(
                parent=self.parents[level - 1],
                image=ImageRef.from_pil(image=pil_image, dpi=72),
                caption=None,
                content_layer=self.content_layer,
            )
        except (UnidentifiedImageError, OSError):
            _log.warning("Warning: DrawingML image cannot be loaded by Pillow")
            doc.add_picture(
                parent=self.parents[level - 1],
                caption=None,
                content_layer=self.content_layer,
            )

        return

    def _add_header_footer(self, docx_obj: DocxDocument, doc: DoclingDocument) -> None:
        """
        添加章节页眉和页脚。

        页眉和页脚添加到装饰内容中，仅解析文本段落。
        段落附加到页眉或页脚的单个组项中。
        如果文档有新页眉和页脚的章节，它们将在新组项中解析。

        Args:
            docx_obj: 要解析的 docx 文档对象。
            doc: 要从 docx_obj 添加页眉和页脚的 DoclingDocument 对象。
        """
        current_layer = self.content_layer
        base_parent = self.parents[0]
        self.content_layer = ContentLayer.FURNITURE
        for sec_idx, section in enumerate(docx_obj.sections):
            if sec_idx > 0 and not section.different_first_page_header_footer:
                continue

            hdr = (
                section.first_page_header
                if section.different_first_page_header_footer
                else section.header
            )
            par = [txt for txt in (par.text.strip() for par in hdr.paragraphs) if txt]
            tables = hdr.tables
            has_blip = self._has_blip(hdr._element)
            if par or tables or has_blip:
                self.parents[0] = doc.add_group(
                    label=GroupLabel.SECTION,
                    name="page header",
                    content_layer=self.content_layer,
                )
                self._walk_linear(hdr._element, doc)

            ftr = (
                section.first_page_footer
                if section.different_first_page_header_footer
                else section.footer
            )
            par = [txt for txt in (par.text.strip() for par in ftr.paragraphs) if txt]
            tables = ftr.tables
            has_blip = self._has_blip(ftr._element)
            if par or tables or has_blip:
                self.parents[0] = doc.add_group(
                    label=GroupLabel.SECTION,
                    name="page footer",
                    content_layer=self.content_layer,
                )
                self._walk_linear(ftr._element, doc)

        self.content_layer = current_layer
        self.parents[0] = base_parent
